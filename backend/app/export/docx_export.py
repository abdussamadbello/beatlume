"""DOCX exporter using python-docx."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.export.base import BaseExporter, ExportOptions, ExportResult


class DOCXExporter(BaseExporter):
    """Generate Word DOCX documents."""

    def export(
        self,
        story: dict,
        chapters: list[dict],
        settings: list[dict],
        options: ExportOptions,
        on_progress: callable | None = None,
    ) -> ExportResult:
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

        title = self._get_setting(settings, "title") or story.get("title", "Untitled")
        author = self._get_setting(settings, "author", "")
        genre = self._get_setting(settings, "genre", "")
        word_count = self._count_words(chapters)

        # Modify default style
        style = doc.styles["Normal"]
        style.font.size = Pt(options.font_size)
        style.paragraph_format.first_line_indent = Inches(0.5)

        # Title page
        if options.include_title_page:
            for _ in range(8):
                doc.add_paragraph("")

            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = tp.add_run(title)
            run.bold = True
            run.font.size = Pt(24)

            if author:
                ap = doc.add_paragraph()
                ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ap.add_run(f"by {author}")
                run.font.size = Pt(14)

            if genre:
                gp = doc.add_paragraph()
                gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                gp.add_run(genre).font.size = Pt(12)

            wp = doc.add_paragraph()
            wp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            wp.add_run(f"{word_count:,} words").font.size = Pt(12)

            doc.add_page_break()

        # Chapters
        total = len(chapters)
        for idx, ch in enumerate(chapters):
            if options.include_chapter_headers:
                ch_title = ch.get("title", f"Chapter {ch.get('num', idx + 1)}")
                header = f"CHAPTER {ch.get('num', idx + 1)}: {ch_title.upper()}"
                h = doc.add_heading(header, level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            content = ch.get("content", "")
            paras = self._split_prose(content, options.include_scene_breaks)

            for p in paras:
                if p == "###":
                    sb = doc.add_paragraph()
                    sb.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sb.add_run("# # #")
                else:
                    doc.add_paragraph(p)

            if idx < total - 1:
                doc.add_page_break()

            if on_progress:
                on_progress((idx + 1) / total)

        buf = io.BytesIO()
        doc.save(buf)
        slug = self._slugify(title)

        return ExportResult(
            file_bytes=buf.getvalue(),
            filename=f"{slug}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            word_count=word_count,
        )
